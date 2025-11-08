from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

def create_editable_docx(md_file, docx_file):
    """Create a fully editable DOCX file from markdown"""
    
    # Remove existing file
    if os.path.exists(docx_file):
        try:
            os.remove(docx_file)
        except:
            pass
    
    # Create new document
    doc = Document()
    
    # Remove any protection by accessing core properties
    try:
        # Force editable document
        core_props = doc.core_properties
    except:
        pass
    
    # Set document properties to editable
    doc.core_properties.title = "Computation Offloading DQN Report"
    doc.core_properties.author = ""
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines at the start
        if not line and i == 0:
            i += 1
            continue
        
        # Main title (# Title)
        if line.startswith('# ') and not line.startswith('##'):
            text = line[2:].strip()
            p = doc.add_heading(text, level=0)
            if text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Heading 1 (## Heading)
        if line.startswith('## ') and not line.startswith('###'):
            text = line[3:].strip()
            if text == 'Project Report':
                p = doc.add_heading(text, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue
        
        # Heading 2 (### Heading)
        if line.startswith('### ') and not line.startswith('####'):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        # Heading 3 (#### Heading)
        if line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Horizontal rule
        if line == '---':
            p = doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # Table detection
        if '|' in line and line.count('|') >= 3:
            table_lines = []
            # Collect table rows
            j = i
            while j < len(lines) and '|' in lines[j] and lines[j].count('|') >= 3:
                if '---' not in lines[j] or lines[j].strip() != '|---|':
                    table_lines.append(lines[j])
                j += 1
            
            if table_lines:
                # Parse headers
                headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
                # Create table
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add header row
                header_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for row_line in table_lines[1:]:
                    row_data = [d.strip() for d in row_line.split('|')[1:-1]]
                    if len(row_data) == len(headers):
                        row_cells = table.add_row().cells
                        for idx, data in enumerate(row_data):
                            row_cells[idx].text = data
                
                i = j
                continue
        
        # Code block
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Bullet list
        if line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            # Handle bold in list
            if '**' in text:
                p.clear()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            if '**' in text:
                p.clear()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            i += 1
            continue
        
        # Regular paragraph with bold
        if line:
            p = doc.add_paragraph()
            if '**' in line:
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                p.add_run(line)
        
        i += 1
    
    # Save document
    try:
        doc.save(docx_file)
        
        # Set file permissions (Windows)
        try:
            os.chmod(docx_file, 0o666)
        except:
            pass
        
        # Remove read-only attribute using Windows command
        try:
            import subprocess
            subprocess.run(['attrib', '-R', docx_file], check=False, capture_output=True)
        except:
            pass
        
        print(f"Created editable DOCX: {docx_file}")
        print("The file should now be fully editable in Microsoft Word.")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

if __name__ == "__main__":
    # Try with a new filename to avoid file lock issues
    output_file = 'Report_Editable.docx'
    if create_editable_docx('Report.md', output_file):
        print(f"\nSUCCESS! Created {output_file}")
        print("Please close Word if it's open, then:")
        print(f"1. Open {output_file} in Word")
        print("2. If you see 'Protected View', click 'Enable Editing'")
        print("3. The document should now be fully editable")

